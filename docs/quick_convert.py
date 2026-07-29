#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""快速 Markdown to Word 转换器"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

print(u"开始处理...")

md_file = r'd:\IdeaProjects\untitled\cloudai-fusion\docs\产品介绍_功能详解.md'
docx_file = r'd:\IdeaProjects\untitled\cloudai-fusion\docs\产品介绍_功能详解.docx'

with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

print("文件已读取")

# 创建文档
doc = Document()
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

lines = content.split('\n')
i = 0

while i < len(lines):
    line = lines[i]
    
    if not line.strip():
        i += 1
        continue
    
    if line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('```'):
        # Skip code block
        j = i + 1
        while j < len(lines) and not lines[j].startswith('```'):
            j += 1
        i = j + 1
        continue
    elif '|' in line and ('---' in ''.join(lines[i:i+5])):
        # Table - skip for now
        while i < len(lines) and '|' in lines[i]:
            i += 1
        continue
    elif line.strip().startswith('- ') and not '|' in line:
        # List item
        items = []
        while i < len(lines) and lines[i].strip().startswith('- '):
            items.append(lines[i].strip()[2:])
            i += 1
        
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(item)
            run.font.size = Pt(10.5)
    else:
        # Regular text
        p = doc.add_paragraph(line.rstrip())
    
    i += 1

print("保存中...")
doc.save(docx_file)
print("完成！文档：" + docx_file)
