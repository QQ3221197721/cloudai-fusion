#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown to Word converter for CloudAI Fusion product introduction
Uses python-docx library for Word document generation
支持中文编码和复杂格式（表格、列表、代码块）
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def parse_markdown(md_file):
    """Parse markdown file content"""
    with open(md_file, 'r', encoding='utf-8') as f:
        return f.read()


def convert_to_word(content, output_file):
    """Convert markdown content to Word document"""
    doc = Document()
    
    # Set default font
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)  # 小四
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[2:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[3:], level=3)
        
        # Table of contents or horizontal rule
        elif line == '---':
            pass  # Just skip
            
        # Tables (lines starting with '|')
        elif line.strip().startswith('|'):
            table_data, end_idx = extract_table_block(lines, i)
            if table_data:
                add_table(doc, table_data)
                i = end_idx
                continue
        
        # Code blocks (starting with ```)
        elif line.startswith('```'):
            code_lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            lang_type = code_lang if code_lang else 'text'
            add_code_block(doc, code_lines, lang_type)
            i += 1  # Skip closing ```
            continue
        
        # List items (starting with '- ')
        elif line.strip().startswith('- '):
            list_items, end_idx = extract_list_block(lines, i)
            add_bulleted_list(doc, list_items)
            i = end_idx
            continue
        
        # Regular text paragraphs
        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Bold first few words if they are keywords
            words = line.rstrip().split()
            if len(words) >= 2 and ':' in words[-1]:
                # This looks like a feature name
                run = p.add_run(words[0] + ' ' + words[1])
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                
                remaining = ' '.join(words[2:])
                if remaining:
                    p.add_run(' ' + remaining)
            else:
                run = p.add_run(line.rstrip())
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(u"转换完成！")
    print("文档已保存至：" + output_file)


def extract_table_block(lines, start_idx):
    """Extract a complete table block from lines"""
    rows = []
    i = start_idx
    
    # Collect all table header/metadata rows
    while i < len(lines):
        line = lines[i]
        
        # Stop if we hit code block or other element
        if line.strip().startswith('**'):
            rows.append(extract_row(line))
            i += 1
            break
        elif line.strip().startswith('| ') or (not line.strip().startswith('- ') and '|' in line):
            rows.append(extract_row(line))
            i += 1
        elif line.strip().startswith('- '):
            # Continue collecting list items under the table context
            break
        else:
            break
    
    if len(rows) >= 2:
        return (rows, i - 1)
    return (None, start_idx)


def extract_row(line):
    """Extract cells from a markdown table row"""
    # Remove leading/trailing |
    if line.strip().startswith('|'):
        parts = line.strip()[1:-1].split('|')
    else:
        parts = line.strip().split('|')
    
    cells = [p.strip() for p in parts if p.strip()]
    return cells


def add_table(doc, table_data):
    """Add a table to the document"""
    if not table_data or len(table_data) < 2:
        return
    
    max_cols = max(len(row) for row in table_data)
    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = 'Table Grid'
    
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            
            run = paragraph.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # Bold first column
            if col_idx == 0:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def extract_list_block(lines, start_idx):
    """Extract a list block from lines"""
    items = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('- '):
            item_text = line.strip()[2:].strip()
            items.append(item_text)
            i += 1
        elif line.strip() and not line.strip().startswith('|'):
            # Non-list, non-table content - stop
            break
        else:
            i += 1
    
    return (items, i - 1)


def add_bulleted_list(doc, items):
    """Add a bulleted list to the document"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_code_block(doc, code_lines, language):
    """Add a code block to the document"""
    # Language identifier
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    lang_run = p.add_run(u"# " + language.upper())
    lang_run.bold = True
    lang_run.font.color.rgb = RGBColor(0, 128, 0)
    
    # Code content
    code_text = '\n'.join(code_lines).strip()
    if code_text:
        code_p = doc.add_paragraph()
        code_p.paragraph_format.space_before = Pt(3)
        code_p.paragraph_format.space_after = Pt(3)
        
        code_run = code_p.add_run(code_text)
        code_run.font.family = 'Consolas'
        code_run.font.size = Pt(9)
        code_run.font.color.rgb = RGBColor(50, 50, 50)


if __name__ == '__main__':
    # Input and output files
    md_file = r'd:\IdeaProjects\untitled\cloudai-fusion\docs\产品介绍_功能详解.md'
    docx_file = r'd:\IdeaProjects\untitled\cloudai-fusion\docs\产品介绍_功能详解.docx'
    
    print(u"开始转换 Markdown 到 Word...")
    print("输入：" + md_file)
    print("输出：" + docx_file)
    print("-" * 60)
    
    try:
        # Parse markdown
        content = parse_markdown(md_file)
        
        # Convert to Word
        convert_to_word(content, docx_file)
    except Exception as e:
        print(u"错误：" + str(e))
        import traceback
        traceback.print_exc()
